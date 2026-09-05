import tempfile
import unittest
from pathlib import Path

from extractor.document_parser_v3 import assess_text_quality, parse_file_v3


class DocumentParserV3Tests(unittest.TestCase):
    def test_quality_gate_rejects_too_short_text(self):
        report = assess_text_quality("姓名 张三\n技能 Python")
        self.assertFalse(report.passed)
        self.assertTrue(report.fallback_required)
        self.assertIn("too_little_text", report.flags)

    def test_quality_gate_accepts_normal_resume_like_text(self):
        text = "\n".join([
            "候选人A 人工智能专业",
            "项目经历：构建图像分类系统，使用PyTorch训练ResNet模型并完成验证。",
            "实习经历：参与数据清洗、模型评估与接口开发，记录实验结果并撰写技术文档。",
            "技能：Python、PyTorch、SQL、Docker。",
        ]) * 2
        report = assess_text_quality(text)
        self.assertTrue(report.passed)

    def test_half_pdf_pages_empty_requires_fallback(self):
        import fitz
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "resume.pdf"
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), "Candidate resume project experience Python PyTorch model training database deployment testing " * 3)
            doc.new_page()
            doc.save(str(path))
            doc.close()
            result = parse_file_v3(path)
            self.assertEqual(result.quality.page_count, 2)
            self.assertEqual(result.quality.empty_page_count, 1)
            self.assertTrue(result.quality.fallback_required)
            self.assertIn("pdf_many_text_empty_pages", result.quality.flags)

    def test_one_blank_trailing_page_out_of_three_does_not_trigger_page_ratio_flag(self):
        import fitz
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "resume.pdf"
            doc = fitz.open()
            for i in range(2):
                page = doc.new_page()
                y = 72
                for line in [
                    f"Candidate {i} project experience and technical work",
                    "Python PyTorch model training evaluation and deployment",
                    "Database API testing documentation collaboration results",
                ]:
                    page.insert_text((72, y), line)
                    y += 20
            doc.new_page()
            doc.save(str(path))
            doc.close()
            result = parse_file_v3(path)
            self.assertAlmostEqual(result.quality.empty_page_ratio, 1/3)
            self.assertNotIn("pdf_many_text_empty_pages", result.quality.flags)

    def test_docx_preserves_paragraph_table_paragraph_order(self):
        from docx import Document
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "resume.docx"
            doc = Document()
            doc.add_paragraph("FIRST_PARAGRAPH")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "TABLE_A"
            table.cell(0, 1).text = "TABLE_B"
            doc.add_paragraph("LAST_PARAGRAPH")
            doc.save(path)
            result = parse_file_v3(path)
            self.assertLess(result.text.index("FIRST_PARAGRAPH"), result.text.index("TABLE_A"))
            self.assertLess(result.text.index("TABLE_A"), result.text.index("LAST_PARAGRAPH"))


if __name__ == "__main__":
    unittest.main()
