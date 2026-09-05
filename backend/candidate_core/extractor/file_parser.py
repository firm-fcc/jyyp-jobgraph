from pathlib import Path


def parse_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8-sig") as f:
        return f.read()


def parse_docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    texts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                texts.append(" ".join(row_text))

    return "\n".join(texts)


def parse_pdf(path: str) -> str:
    import fitz

    doc = fitz.open(path)
    texts = []

    for page in doc:
        page_text = page.get_text()
        if page_text.strip():
            texts.append(page_text)

    return "\n".join(texts)


def parse_file(path: str) -> str:
    suffix = Path(path).suffix.lower()

    if suffix == ".txt":
        return parse_txt(path)

    if suffix == ".docx":
        return parse_docx(path)

    if suffix == ".pdf":
        return parse_pdf(path)

    raise ValueError(f"暂不支持的文件类型：{suffix}")
